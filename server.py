import os
import re
from pathlib import Path
from datetime import datetime
from mcp.server.fastmcp import FastMCP

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Config ---
BASE_DIR = Path.home() / "Desktop" / "Job Applications"

# Initialize MCP server
mcp = FastMCP("Job Application Agent")


# --- Helper ---
def sanitize(name: str) -> str:
    """Remove characters that are unsafe in folder names."""
    return "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip()


# --- Docx helpers ---

def set_paragraph_border_bottom(paragraph, color="CCCCCC", size=6):
    """Add a bottom border to a paragraph (used as a section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_horizontal_rule(doc, color="AAAAAA"):
    """Add a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    set_paragraph_border_bottom(p, color=color, size=4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins in inches."""
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


# --- Resume section detection ---

# Section header keywords commonly found in resumes
RESUME_SECTION_KEYWORDS = {
    "summary", "objective", "profile", "about",
    "experience", "work experience", "professional experience", "employment",
    "education", "academic", "academics",
    "skills", "technical skills", "core competencies", "competencies",
    "projects", "portfolio",
    "certifications", "certificates", "licenses",
    "awards", "honors", "achievements",
    "publications", "volunteer", "volunteering",
    "languages", "interests", "activities",
    "references",
}

def is_section_header(line: str) -> bool:
    """Return True if this line looks like a resume section header."""
    stripped = line.strip().rstrip(":").lower()
    if stripped in RESUME_SECTION_KEYWORDS:
        return True
    # All-caps short line is likely a header too
    if line.strip().isupper() and 2 <= len(line.strip().split()) <= 5:
        return True
    return False

def is_bullet(line: str) -> tuple[bool, str]:
    """Return (True, text) if line is a bullet point, else (False, line)."""
    # Matches: • - * – — or leading digits like "1." "2)"
    m = re.match(r"^[\s]*[•\-\*–—]\s+(.*)", line)
    if m:
        return True, m.group(1).strip()
    m = re.match(r"^[\s]*\d+[.)]\s+(.*)", line)
    if m:
        return True, m.group(1).strip()
    return False, line


# --- Build resume .docx ---

def build_resume_docx(resume_text: str, output_path: Path) -> None:
    """
    Parse plain-text resume and produce a formatted Word document.
    Detects: name (first non-blank line), contact info, section headers,
    bullets, and body paragraphs. Mirrors the structure of the source resume.
    """
    doc = Document()
    set_doc_margins(doc, top=0.85, bottom=0.85, left=1.0, right=1.0)

    # --- Base font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = resume_text.strip().splitlines()
    if not lines:
        doc.add_paragraph("(empty resume)")
        doc.save(str(output_path))
        return

    # --- Pass 1: identify structure ---
    # First non-blank line → candidate name
    name_line = ""
    start_idx = 0
    for i, line in enumerate(lines):
        if line.strip():
            name_line = line.strip()
            start_idx = i + 1
            break

    # Collect contact/header lines: lines right after the name until
    # we hit a blank line or a section header
    contact_lines = []
    body_start = start_idx
    for i in range(start_idx, len(lines)):
        l = lines[i].strip()
        if not l:
            body_start = i + 1
            break
        if is_section_header(l):
            body_start = i
            break
        contact_lines.append(l)
    else:
        body_start = len(lines)

    # --- Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name_line)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    name_para.paragraph_format.space_after = Pt(2)

    # --- Contact line(s) ---
    if contact_lines:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_lines))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        contact_para.paragraph_format.space_after = Pt(4)

    # Divider under header
    add_horizontal_rule(doc, color="1F497D")

    # --- Body ---
    i = body_start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Section header
        if is_section_header(stripped):
            sec_para = doc.add_paragraph()
            sec_para.paragraph_format.space_before = Pt(10)
            sec_para.paragraph_format.space_after = Pt(2)
            sec_run = sec_para.add_run(stripped.rstrip(":").upper())
            sec_run.bold = True
            sec_run.font.size = Pt(11)
            sec_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            set_paragraph_border_bottom(sec_para, color="1F497D", size=4)
            i += 1
            continue

        # Bullet
        is_b, bullet_text = is_bullet(line)
        if is_b:
            b_para = doc.add_paragraph(style="List Bullet")
            b_para.paragraph_format.space_before = Pt(1)
            b_para.paragraph_format.space_after = Pt(1)
            # Left indent + hanging to match typical resume bullet
            b_para.paragraph_format.left_indent = Inches(0.25)
            run = b_para.add_run(bullet_text)
            run.font.size = Pt(10.5)
            i += 1
            continue

        # Role/company lines: bold if short and likely a title
        # Heuristic: non-bullet, non-section, ≤ 10 words, no terminal period → bold
        words = stripped.split()
        looks_like_title = (
            len(words) <= 10
            and not stripped.endswith(".")
            and not stripped.endswith(",")
            and i + 1 < len(lines)
            and (not lines[i + 1].strip() or is_section_header(lines[i + 1].strip()) or is_bullet(lines[i + 1])[0] or (lines[i + 1].strip() and len(lines[i + 1].strip().split()) <= 10))
        )

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(stripped)
        r.font.size = Pt(10.5)
        if looks_like_title:
            r.bold = True
        i += 1

    doc.save(str(output_path))


# --- Build cover letter .docx ---

def build_cover_letter_docx(cover_letter_text: str, output_path: Path) -> None:
    """
    Format the cover letter as a clean Word document.
    Preserves paragraph breaks, applies professional typography.
    """
    doc = Document()
    set_doc_margins(doc, top=1.0, bottom=1.0, left=1.15, right=1.15)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    paragraphs = cover_letter_text.strip().split("\n\n")

    for idx, para_block in enumerate(paragraphs):
        lines_in_block = [l.strip() for l in para_block.strip().splitlines() if l.strip()]
        if not lines_in_block:
            continue

        text = " ".join(lines_in_block)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(11)

        # First para (salutation) or last para (sign-off) — left aligned, normal weight
        # Body paragraphs — justified
        if idx == 0 or idx == len(paragraphs) - 1:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(str(output_path))


# --- Tools ---

@mcp.tool()
def save_application(
    company: str,
    role: str,
    job_id: str,
    resume_content: str,
    job_description: str,
    cover_letter: str
) -> str:
    """
    Save a job application to the local filesystem.
    Call this when the user says 'save now'.
    Creates a folder structure: Desktop/Job Applications/Company/Role-JobID/
    Saves resume and cover letter as formatted Word (.docx) files.
    Saves job description as a plain text file.

    Args:
        company: Company name e.g. Amazon
        role: Job title e.g. Senior Product Manager
        job_id: Job ID from the posting e.g. JR-123456
        resume_content: The full customized resume text
        job_description: The full job description text
        cover_letter: The full cover letter text
    """
    try:
        # Validate required fields
        if not sanitize(company):
            return "Error: Company name is missing or invalid. Please provide a valid company name."
        if not sanitize(role):
            return "Error: Role title is missing or invalid. Please provide a valid role title."
        if not sanitize(job_id):
            return "Error: Job ID is missing or invalid. Please provide a valid job ID."
        if not resume_content.strip():
            return "Error: Resume content is empty. Please provide the resume before saving."
        if not job_description.strip():
            return "Error: Job description is empty. Please provide the job description before saving."
        if not cover_letter.strip():
            return "Error: Cover letter is empty. Please write a cover letter before saving."

        # Build folder path
        safe_company = sanitize(company)
        safe_role = sanitize(role)
        safe_job_id = sanitize(job_id)
        folder_name = f"{safe_role}-{safe_job_id}"
        folder_path = BASE_DIR / safe_company / folder_name

        # Create folders
        folder_path.mkdir(parents=True, exist_ok=True)

        # Save job description as plain text (human-readable reference)
        (folder_path / "job_description.txt").write_text(job_description, encoding="utf-8")

        # Save resume as formatted Word document
        resume_path = folder_path / "resume.docx"
        build_resume_docx(resume_content, resume_path)

        # Save cover letter as formatted Word document
        cover_letter_path = folder_path / "cover_letter.docx"
        build_cover_letter_docx(cover_letter, cover_letter_path)

        return (
            f"Saved successfully!\n"
            f"Location: {folder_path}\n"
            f"Files saved:\n"
            f"  - job_description.txt\n"
            f"  - resume.docx  (formatted Word document)\n"
            f"  - cover_letter.docx  (formatted Word document)"
        )

    except Exception as e:
        return f"Error saving application: {str(e)}"


@mcp.tool()
def list_applications() -> str:
    """
    List all job applications saved on the desktop.
    Shows company, role, job ID and date saved.
    """
    try:
        if not BASE_DIR.exists():
            return "No applications found. No Job Applications folder exists yet."

        results = []
        for company_folder in sorted(BASE_DIR.iterdir()):
            if not company_folder.is_dir():
                continue
            for job_folder in sorted(company_folder.iterdir()):
                if not job_folder.is_dir():
                    continue
                created = datetime.fromtimestamp(
                    job_folder.stat().st_birthtime
                ).strftime("%Y-%m-%d")
                results.append(f"  {company_folder.name} | {job_folder.name} | saved {created}")

        if not results:
            return "No applications saved yet."

        return "Your saved applications:\n" + "\n".join(results)

    except Exception as e:
        return f"Error listing applications: {str(e)}"


# --- Prompts ---

@mcp.prompt()
def review_resume() -> str:
    """Review resume against the job description."""
    return (
        "Before proceeding, check that both a job description AND a resume have been "
        "provided in this conversation. If either is missing, stop and ask the user to "
        "paste the missing item before continuing.\n\n"
        "If both are present: please review my resume against the job description. "
        "Identify the top 5 gaps or mismatches, and suggest specific improvements "
        "to better align my resume with this role. Be direct and specific."
    )


@mcp.prompt()
def rewrite_resume() -> str:
    """Rewrite resume tailored to the job description."""
    return (
        "Before proceeding, check that both a job description AND a resume have been "
        "provided in this conversation. If either is missing, stop and ask the user to "
        "paste the missing item before continuing.\n\n"
        "If both are present: please rewrite my resume tailored specifically to this "
        "job description. You must follow these rules strictly:\n"
        "- Use ONLY the experience, skills and roles already present in my resume\n"
        "- Do NOT add, invent or infer anything that is not already in my resume\n"
        "- Do NOT inflate or exaggerate any experience or skill\n"
        "- Do NOT add keywords that do not reflect my actual experience\n"
        "- Your job is to reframe and reorder what is already there — not to fabricate\n"
        "- Preserve my voice and keep it authentic to who I am\n\n"
        "The goal is an honest, tailored resume — not a keyword-optimized one."
    )


@mcp.prompt()
def write_cover_letter() -> str:
    """Write a cover letter for this role."""
    return (
        "Before proceeding, check that both a job description AND a resume have been "
        "provided in this conversation. If either is missing, stop and ask the user to "
        "paste the missing item before continuing.\n\n"
        "If both are present: please write a compelling cover letter for this role. "
        "You must follow these rules strictly:\n"
        "- Base the cover letter ONLY on experience and skills present in my resume\n"
        "- Do NOT invent, exaggerate or infer anything not already in my resume\n"
        "- Do NOT make claims about my experience that are not supported by my resume\n\n"
        "Structure: 3 paragraphs — why this company, why I am a genuine fit based on "
        "my actual experience, and a clear call to action. Keep it concise and authentic."
    )


@mcp.prompt()
def save_now() -> str:
    """Trigger to save the application artifacts."""
    return (
        "Before saving, check that all of the following are present in our conversation:\n"
        "  1. Job description\n"
        "  2. Customized resume\n"
        "  3. Cover letter\n\n"
        "If any are missing, stop and ask the user to provide them before saving.\n\n"
        "If all are present:\n"
        "  - Extract the company name and role title from the job description\n"
        "  - Look for a Job ID in the job description (e.g. JR-123456, REQ-789, #12345)\n"
        "  - If a Job ID is found, confirm it with the user before saving\n"
        "  - If no Job ID is found, ask the user: 'I could not find a Job ID in the job "
        "description. Could you provide it? If there is no Job ID just type NONE and I "
        "will use today's date instead.'\n\n"
        "Once confirmed, call save_application with the company, role, job ID "
        "(or today's date in YYYY-MM-DD format if NONE), and all three artifacts."
    )


# --- Run ---
if __name__ == "__main__":
    mcp.run()
